import asyncio
import hashlib
import json
import os
import re
import shutil
import tempfile
import time as time_module
import uuid
import zipfile
from collections import Counter
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from openai import OpenAI
try:
    from docxtpl import DocxTemplate
except Exception:
    DocxTemplate = None

from config import DEEPSEEK_API_KEY, DEEPSEEK_MODEL, OPENAI_API_KEY, WORD_AI_PROVIDER, WORD_OPENAI_MODEL, logger
from database import (
    add_to_conversation,
    delete_dialog_file_state,
    get_dialog_file_state,
    upsert_dialog_file_state,
)
from group_utils import get_dialog_key

DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
openai_word_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

SUPPORTED_WORD_EXTENSIONS = (".docx",)
UNSUPPORTED_WORD_EXTENSIONS = (".doc", ".rtf", ".odt")


class WordProcessingError(Exception):
    pass


def is_word_file(file_name: str) -> bool:
    return (file_name or "").lower().endswith(SUPPORTED_WORD_EXTENSIONS + UNSUPPORTED_WORD_EXTENSIONS)


def is_supported_word_file(file_name: str) -> bool:
    return (file_name or "").lower().endswith(SUPPORTED_WORD_EXTENSIONS)


def looks_like_word_edit_request(text: str) -> bool:
    if not text:
        return False
    lowered = text.lower()
    edit_words = (
        "измени", "исправь", "поправь", "отредактируй", "добавь", "удали", "замени",
        "вставь", "перепиши", "дополни", "сократи", "расширь", "сформулируй в договоре",
        "добавить", "изменить", "правки", "правку", "новую редакцию", "этапность оплат",
        "порядок оплаты", "порядок расчетов", "раздел", "пункт", "пункты", "приложение",
    )
    file_words = ("word", "docx", "ворд", "документ", "договор", "акт", "письмо", "коммерческое")
    if any(word in lowered for word in edit_words):
        return True
    return "сделай" in lowered and any(word in lowered for word in file_words)


def is_create_word_request(text: str) -> bool:
    if not text:
        return False
    lowered = text.lower()
    create_words = ("создай", "сделай", "сформируй", "подготовь", "собери")
    if not any(word in lowered for word in create_words):
        return False
    # Не перехватываем общие управленческие задачи вида «подготовь РПЗ».
    return any(word in lowered for word in ("word", "docx", "ворд", "договор", "акт", "письмо", "протокол"))


def _safe_filename(name: str, default: str = "document_result.docx") -> str:
    name = (name or default).strip().replace("\\", "_").replace("/", "_")
    name = re.sub(r"[^\w\-.а-яА-ЯёЁ ]+", "_", name, flags=re.UNICODE).strip(" ._")
    if not name:
        name = default
    if not name.lower().endswith(".docx"):
        name += ".docx"
    return name[:120]


def _delivery_filename(path: str) -> str:
    name = os.path.basename(path or "document_result.docx")
    return _safe_filename(re.sub(r"^[0-9a-f]{6}_", "", name, flags=re.IGNORECASE))


def _word_cache_dir() -> str:
    """Папка для последнего Word-файла пользователя.

    На Railway можно задать WORD_STORAGE_DIR=/data/word_cache.
    Если не задано, но есть DB_FILE=/data/bot_data.db, файлы кладутся рядом с базой.
    """
    explicit = os.getenv("WORD_STORAGE_DIR")
    if explicit:
        base = explicit
    else:
        db_file = os.getenv("DB_FILE")
        if db_file:
            base = os.path.join(os.path.dirname(os.path.abspath(db_file)), "word_cache")
        elif os.getenv("RAILWAY_VOLUME_MOUNT_PATH"):
            base = os.path.join(os.getenv("RAILWAY_VOLUME_MOUNT_PATH"), "word_cache")
        else:
            base = os.path.join(tempfile.gettempdir(), "mitra_word_cache")
    os.makedirs(base, exist_ok=True)
    return base


def _word_template_dir() -> str:
    explicit = os.getenv("WORD_TEMPLATE_DIR")
    if explicit:
        base = explicit
    else:
        db_file = os.getenv("DB_FILE")
        if db_file:
            base = os.path.join(os.path.dirname(os.path.abspath(db_file)), "word_templates")
        elif os.getenv("RAILWAY_VOLUME_MOUNT_PATH"):
            base = os.path.join(os.getenv("RAILWAY_VOLUME_MOUNT_PATH"), "word_templates")
        else:
            base = os.path.join(_word_cache_dir(), "templates")
    os.makedirs(base, exist_ok=True)
    return base


def _list_word_templates() -> List[Dict[str, str]]:
    template_dir = _word_template_dir()
    templates = []
    for path in sorted(Path(template_dir).glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        templates.append({"name": path.stem, "file_name": path.name, "path": str(path)})
    return templates


def _looks_like_template_upload(text: str) -> bool:
    lowered = (text or "").lower()
    return "шаблон" in lowered and any(word in lowered for word in ("word", "docx", "договор", "акт", "кп", "письмо", "коммерческ"))


def _looks_like_template_list_request(text: str) -> bool:
    lowered = (text or "").lower()
    return "шаблон" in lowered and any(word in lowered for word in ("какие", "список", "покажи", "есть"))


def remember_word_template(source_path: str, file_name: str, caption: str = "") -> Dict[str, str]:
    safe_name = _safe_filename(file_name or "word_template.docx")
    stem = Path(safe_name).stem
    lowered = (caption or "").lower()
    for marker in ("шаблон договора", "шаблон договор"):
        if marker in lowered and "договор" not in stem.lower():
            stem = "шаблон_договора"
            break
    if "шаблон кп" in lowered and "кп" not in stem.lower():
        stem = "шаблон_кп"
    if "шаблон акта" in lowered and "акт" not in stem.lower():
        stem = "шаблон_акта"
    if "шаблон письма" in lowered and "письм" not in stem.lower():
        stem = "шаблон_письма"

    target_name = _safe_filename(stem + ".docx")
    target_path = os.path.join(_word_template_dir(), target_name)
    shutil.copy2(source_path, target_path)
    return {"name": Path(target_name).stem, "file_name": target_name, "path": target_path}


def _word_template_variables(template_path: str) -> List[str]:
    if not DocxTemplate:
        raise WordProcessingError("Не установлена библиотека docxtpl. Добавьте docxtpl в requirements.txt и перезапустите Railway.")
    tpl = DocxTemplate(template_path)
    try:
        variables = tpl.get_undeclared_template_variables()
    except TypeError:
        variables = tpl.get_undeclared_template_variables({})
    return sorted(str(v) for v in variables)


def _score_template_for_request(template: Dict[str, str], request_text: str) -> int:
    low = (request_text or "").lower()
    name = (template.get("name") or "").lower()
    score = 0
    aliases = {
        "договор": ("договор", "контракт"),
        "кп": ("кп", "коммерческ", "предложен"),
        "акт": ("акт", "прием", "приём", "сдач"),
        "письм": ("письм", "ответ", "заказчик", "клиент"),
        "протокол": ("протокол", "совещан", "встреч"),
    }
    for token in re.split(r"[\s_\-.]+", name):
        if token and token in low:
            score += 4
    for key, words in aliases.items():
        if key in name and any(word in low for word in words):
            score += 10
    return score


def _choose_word_template(request_text: str) -> Optional[Dict[str, str]]:
    templates = _list_word_templates()
    if not templates:
        return None
    scored = sorted(
        ((template, _score_template_for_request(template, request_text)) for template in templates),
        key=lambda item: item[1],
        reverse=True,
    )
    if scored[0][1] <= 0 and len(templates) != 1:
        return None
    return scored[0][0]


def _render_word_template(template_path: str, context: Dict[str, Any], output_path: str) -> None:
    if not DocxTemplate:
        raise WordProcessingError("Не установлена библиотека docxtpl. Добавьте docxtpl в requirements.txt и перезапустите Railway.")
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    tpl.save(output_path)


def remember_word_file(context, chat_id: int, source_path: str, file_name: str, dialog_key: Optional[str] = None) -> Dict[str, Any]:
    dialog_key = str(dialog_key or chat_id)
    word_files = context.user_data.setdefault("word_files", {}) if hasattr(context, "user_data") else {}
    old = word_files.get(dialog_key)
    if old and old.get("path") and os.path.exists(old["path"]):
        try:
            os.unlink(old["path"])
        except OSError:
            pass

    safe_name = _safe_filename(file_name or "document.docx")
    stored_name = f"chat_{chat_id}_{uuid.uuid4().hex[:8]}_{safe_name}"
    stored_path = os.path.join(_word_cache_dir(), stored_name)
    shutil.copy2(source_path, stored_path)

    info = {"path": stored_path, "file_name": safe_name, "saved_at": time_module.time()}
    word_files[dialog_key] = info
    upsert_dialog_file_state(
        dialog_key,
        "word",
        safe_name,
        stored_path,
        info["saved_at"],
    )
    context.user_data["awaiting_word_request_by_dialog"] = {
        **context.user_data.get("awaiting_word_request_by_dialog", {}),
        dialog_key: True,
    }
    return info


def _get_recent_word_context(context, dialog_key: str, max_age_seconds: int = 24 * 60 * 60) -> Optional[Dict[str, Any]]:
    dialog_key = str(dialog_key)
    word_files = context.user_data.get("word_files", {}) if hasattr(context, "user_data") else {}
    info = word_files.get(dialog_key)
    if not info:
        info = get_dialog_file_state(dialog_key, "word")
        if info:
            word_files[dialog_key] = info
    if not info:
        return None
    path = info.get("path")
    if not path or not os.path.exists(path):
        word_files.pop(dialog_key, None)
        context.user_data.get("awaiting_word_request_by_dialog", {}).pop(dialog_key, None)
        delete_dialog_file_state(dialog_key, "word")
        return None
    if time_module.time() - float(info.get("saved_at") or 0) > max_age_seconds:
        word_files.pop(dialog_key, None)
        context.user_data.get("awaiting_word_request_by_dialog", {}).pop(dialog_key, None)
        delete_dialog_file_state(dialog_key, "word")
        return None
    return info


def _get_reference_documents(context, dialog_key: str, max_docs: int = 5) -> List[Dict[str, Any]]:
    if not hasattr(context, "user_data"):
        return []
    refs_by_dialog = context.user_data.get("reference_documents_by_dialog", {})
    items = refs_by_dialog.get(str(dialog_key), [])
    return list(items[-max_docs:])


def _remember_reference_document(context, dialog_key: str, file_name: str, text: str, max_docs: int = 8) -> None:
    if not text or not hasattr(context, "user_data"):
        return
    refs_by_dialog = context.user_data.setdefault("reference_documents_by_dialog", {})
    items = refs_by_dialog.setdefault(str(dialog_key), [])
    items.append(
        {
            "file_name": file_name or "document.docx",
            "text": text[:12000],
            "saved_at": time_module.time(),
        }
    )
    refs_by_dialog[str(dialog_key)] = items[-max_docs:]


def _clip_text_at_boundary(text: str, limit: int, suffix: str = "\n... [текст обрезан]") -> str:
    if len(text) <= limit:
        return text
    cut = text[:limit]
    boundary = max(cut.rfind("\n"), cut.rfind(". "), cut.rfind("; "), cut.rfind(", "), cut.rfind(" "))
    if boundary > max(0, limit - 300):
        cut = cut[:boundary].rstrip()
    return cut.rstrip() + suffix


def _format_reference_documents(refs: List[Dict[str, Any]], limit: int = 22000) -> str:
    if not refs:
        return ""
    parts: List[str] = []
    used = 0
    for idx, ref in enumerate(refs, start=1):
        name = str(ref.get("file_name") or f"reference_{idx}")
        text = str(ref.get("text") or "").strip()
        if not text:
            continue
        header = f"\n=== Reference document {idx}: {name} ===\n"
        remaining = limit - used - len(header)
        if remaining <= 0:
            break
        chunk = _clip_text_at_boundary(text, remaining, "\n... [часть опорного документа скрыта из-за лимита контекста]")
        parts.append(header + chunk)
        used += len(header) + len(chunk)
        if used >= limit:
            break
    return "\n".join(parts).strip()


def _is_word_followup_request(context, text: str, dialog_key: str) -> bool:
    if not text:
        return False
    lowered = text.lower()
    awaiting = context.user_data.get("awaiting_word_request_by_dialog", {}) if hasattr(context, "user_data") else {}
    if awaiting.get(str(dialog_key)):
        return True
    if looks_like_word_edit_request(text):
        return True
    return any(word in lowered for word in ("word", "docx", "ворд", "договор", "пункт", "раздел", "оплат", "расчет"))


def _trim_text(text: str, limit: int = 90000) -> str:
    if len(text) <= limit:
        return text
    return _clip_text_at_boundary(text, limit, "\n... [текст документа обрезан из-за размера файла]")


def _length_cm(value) -> str:
    if value is None:
        return "auto"
    try:
        return f"{value.cm:.2f}"
    except Exception:
        return str(value)


def _document_style_profile(doc: Document) -> str:
    styles = Counter()
    fonts = Counter()
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else "без стиля"
        if (paragraph.text or "").strip():
            styles[style_name] += 1
        for run in paragraph.runs:
            if run.text and run.font.name:
                fonts[run.font.name] += len(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    style_name = paragraph.style.name if paragraph.style is not None else "без стиля"
                    if (paragraph.text or "").strip():
                        styles[style_name] += 1
                    for run in paragraph.runs:
                        if run.text and run.font.name:
                            fonts[run.font.name] += len(run.text)

    style_text = ", ".join(f"{name}: {count}" for name, count in styles.most_common(8)) or "нет данных"
    font_text = ", ".join(name for name, _ in fonts.most_common(5)) or "задаются стилями документа"
    section_lines = []
    for idx, section in enumerate(doc.sections, start=1):
        section_lines.append(
            f"S{idx}: page={_length_cm(section.page_width)}x{_length_cm(section.page_height)} cm; "
            f"margins={_length_cm(section.top_margin)}/{_length_cm(section.right_margin)}/"
            f"{_length_cm(section.bottom_margin)}/{_length_cm(section.left_margin)} cm; "
            f"header={_length_cm(section.header_distance)} cm; footer={_length_cm(section.footer_distance)} cm"
        )
    return (
        "[DOCUMENT PROFILE]\n"
        f"Sections: {len(doc.sections)}; body paragraphs: {len(doc.paragraphs)}; tables: {len(doc.tables)}\n"
        f"Common paragraph styles: {style_text}\n"
        f"Explicit fonts: {font_text}\n"
        + "\n".join(section_lines)
    )


def _iter_body_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _paragraph_format_hint(paragraph: Paragraph) -> str:
    values = []
    if paragraph.alignment is not None:
        values.append(f"align={paragraph.alignment}")
    fmt = paragraph.paragraph_format
    for label, value in (
        ("left", fmt.left_indent),
        ("right", fmt.right_indent),
        ("first", fmt.first_line_indent),
        ("before", fmt.space_before),
        ("after", fmt.space_after),
    ):
        if value is not None:
            values.append(f"{label}={_length_cm(value)}cm")
    if fmt.line_spacing is not None:
        values.append(f"line={fmt.line_spacing}")
    try:
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            values.append("numbered=yes")
    except Exception:
        pass
    ref_run = next((run for run in paragraph.runs if run.text), None)
    if ref_run is not None:
        if ref_run.font.name:
            values.append(f"font={ref_run.font.name}")
        if ref_run.font.size:
            values.append(f"size={ref_run.font.size.pt:g}pt")
        if ref_run.bold:
            values.append("bold=yes")
        if ref_run.italic:
            values.append("italic=yes")
    return ";".join(values) or "inherited"


def document_to_text(path: str, max_paragraphs: int = 500, max_tables: int = 30, max_table_rows: int = 150, max_cell_chars: int = 1400) -> str:
    """Structured DOCX view in original block order with stable target IDs."""
    doc = Document(path)
    parts: List[str] = [_document_style_profile(doc), "\n[DOCUMENT CONTENT]"]
    paragraph_count = 0
    table_count = 0
    truncated_paragraphs = False
    for block in _iter_body_blocks(doc):
        if isinstance(block, Paragraph):
            paragraph_count += 1
            if paragraph_count > max_paragraphs:
                truncated_paragraphs = True
                continue
            text = (block.text or "").strip()
            if text:
                style = block.style.name if block.style is not None else ""
                parts.append(
                    f"[P{paragraph_count}; style={style}; fmt={_paragraph_format_hint(block)}] {text}"
                )
            continue

        table_count += 1
        if table_count > max_tables:
            continue
        table = block
        style = table.style.name if table.style is not None else ""
        parts.append(
            f"\n[T{table_count}; style={style}; rows={len(table.rows)}; cols={len(table.columns)}]"
        )
        seen_cells: Dict[Any, str] = {}
        for r_idx, row in enumerate(table.rows[:max_table_rows], start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                cell_key = cell._tc
                cell_id = f"T{table_count}.R{r_idx}.C{c_idx}"
                if cell_key in seen_cells:
                    parts.append(f"[{cell_id}; merged-with={seen_cells[cell_key]}]")
                    continue
                seen_cells[cell_key] = cell_id
                nonempty = False
                for p_idx, paragraph in enumerate(cell.paragraphs, start=1):
                    text = (paragraph.text or "").strip()
                    if not text:
                        continue
                    nonempty = True
                    style_name = paragraph.style.name if paragraph.style is not None else ""
                    clipped = _clip_text_at_boundary(
                        text,
                        max_cell_chars,
                        " ... [текст ячейки сокращен из-за лимита]",
                    )
                    parts.append(
                        f"[{cell_id}.P{p_idx}; style={style_name}; "
                        f"fmt={_paragraph_format_hint(paragraph)}] {clipped}"
                    )
                if not nonempty:
                    parts.append(f"[{cell_id}.P1; empty]")
        if len(table.rows) > max_table_rows:
            parts.append(f"[T{table_count}] показаны первые {max_table_rows} строк")

    if truncated_paragraphs:
        parts.append(f"[BODY] показаны первые {max_paragraphs} абзацев")
    if len(doc.tables) > max_tables:
        parts.append(f"[TABLES] показаны первые {max_tables} таблиц из {len(doc.tables)}")
    if len(parts) <= 2:
        parts.append("[Документ не содержит извлекаемого текста или состоит из изображений/сканов]")
    return _trim_text("\n".join(parts))


def _call_deepseek_for_word(system_prompt: str, user_prompt: str, max_tokens: int = 3000, temperature: float = 0.1) -> str:
    if not DEEPSEEK_API_KEY:
        raise WordProcessingError("Не настроен DEEPSEEK_API_KEY")

    try:
        word_token_limit = int(os.getenv("WORD_MAX_RESPONSE_TOKENS", "8000"))
    except ValueError:
        word_token_limit = 8000

    data = {
        "model": DEEPSEEK_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "max_tokens": max(512, min(word_token_limit, max_tokens)),
        "temperature": temperature,
    }
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=140)
        response.raise_for_status()
        payload = response.json()
        choice = payload["choices"][0]
        message = choice.get("message") or {}
        content = message.get("content") or ""
        finish_reason = choice.get("finish_reason") or "unknown"
        if not content.strip():
            raise WordProcessingError(f"ИИ вернул пустой ответ (finish_reason={finish_reason})")
        if finish_reason == "length" and not content.strip().endswith("}"):
            raise WordProcessingError("ИИ не успел завершить JSON-план правок (finish_reason=length)")
        return content
    except requests.HTTPError as e:
        detail = (getattr(e.response, "text", "") or "")[:500]
        status = getattr(e.response, "status_code", "unknown")
        raise WordProcessingError(f"ИИ-сервис вернул ошибку HTTP {status}. {detail}") from e
    except requests.RequestException as e:
        raise WordProcessingError(f"ИИ-сервис временно недоступен или не ответил: {e}") from e
    except (KeyError, IndexError, ValueError) as e:
        raise WordProcessingError("ИИ-сервис вернул ответ в неожиданном формате") from e


def _call_openai_for_word(system_prompt: str, user_prompt: str, max_tokens: int = 3000, temperature: float = 0.1) -> str:
    if not openai_word_client:
        raise WordProcessingError("Не настроен OPENAI_API_KEY для обработки Word")

    try:
        word_token_limit = int(os.getenv("WORD_MAX_RESPONSE_TOKENS", "8000"))
    except ValueError:
        word_token_limit = 8000

    try:
        request_args = dict(
            model=WORD_OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            max_tokens=max(512, min(word_token_limit, max_tokens)),
            temperature=temperature,
        )
        wants_json = "json" in system_prompt.lower() and (
            "валидн" in system_prompt.lower() or "only" in system_prompt.lower()
        )
        if wants_json:
            request_args["response_format"] = {"type": "json_object"}
        try:
            response = openai_word_client.chat.completions.create(**request_args)
        except Exception as response_format_error:
            error_text = str(response_format_error).lower()
            if not wants_json or not any(
                marker in error_text
                for marker in ("response_format", "json_object", "unsupported", "not supported")
            ):
                raise
            request_args.pop("response_format", None)
            response = openai_word_client.chat.completions.create(**request_args)
        choice = response.choices[0]
        content = choice.message.content or ""
        finish_reason = choice.finish_reason or "unknown"
        if not content.strip():
            raise WordProcessingError(f"OpenAI вернул пустой ответ (finish_reason={finish_reason})")
        if finish_reason == "length" and not content.strip().endswith("}"):
            raise WordProcessingError("OpenAI не успел завершить JSON-план правок (finish_reason=length)")
        return content
    except WordProcessingError:
        raise
    except Exception as e:
        raise WordProcessingError(f"OpenAI не смог обработать Word-запрос: {e}") from e


def _call_ai_for_word(system_prompt: str, user_prompt: str, max_tokens: int = 3000, temperature: float = 0.1) -> str:
    if WORD_AI_PROVIDER == "openai":
        return _call_openai_for_word(system_prompt, user_prompt, max_tokens=max_tokens, temperature=temperature)
    if WORD_AI_PROVIDER == "auto" and openai_word_client:
        try:
            return _call_openai_for_word(system_prompt, user_prompt, max_tokens=max_tokens, temperature=temperature)
        except WordProcessingError as openai_error:
            logger.warning(f"OpenAI Word fallback failed, trying DeepSeek: {openai_error}")
    return _call_deepseek_for_word(system_prompt, user_prompt, max_tokens=max_tokens, temperature=temperature)


def _extract_json(text: str) -> Dict[str, Any]:
    if not text:
        raise WordProcessingError("ИИ вернул пустой ответ")
    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(cleaned[start:end + 1])
        raise WordProcessingError("ИИ не смог сформировать корректный JSON-план правок") from None


def analyze_word_with_ai(path: str, file_name: str, question: str, user_id: str, reference_context: str = "") -> str:
    text = document_to_text(path)
    question = (question or "Кратко опиши, что находится в Word-документе.").strip()
    system_prompt = (
        "Ты — ассистент, который анализирует Word/DOCX документы. Отвечай на русском. "
        "Используй текст Word-документа и дополнительные документы-основания, если они переданы. "
        "Если часть документа не видна или это скан, честно скажи об этом. "
        "Служебные пометки «текст сокращен из-за лимита» означают только ограничение переданного контекста; "
        "они не доказывают повреждение исходного DOCX и не являются обрывом текста в самом файле."
    )
    user_prompt = (
        f"Файл: {file_name}\n\n"
        f"Текст DOCX:\n{text}\n\n"
        f"Дополнительные документы/контекст (ТЗ, КП, PDF, TXT), если были переданы:\n{reference_context or '[нет]'}\n\n"
        f"Вопрос пользователя:\n{question}"
    )
    answer = _call_ai_for_word(system_prompt, user_prompt, max_tokens=3000, temperature=0.2)
    add_to_conversation(user_id, "user", f"[Word {file_name}] {question}\n{text[:5000]}")
    add_to_conversation(user_id, "assistant", answer)
    return answer


def build_word_patch_with_ai(path: str, file_name: str, request_text: str, reference_context: str = "") -> Dict[str, Any]:
    document_text = document_to_text(
        path,
        max_paragraphs=500,
        max_tables=30,
        max_table_rows=150,
        max_cell_chars=1400,
    )
    reference_context = _clip_text_at_boundary(
        reference_context or "",
        14000,
        "\n... [контекст ТЗ/КП сокращен из-за лимита]",
    )
    system_prompt = """
Ты преобразуешь просьбу пользователя о правке Word/DOCX в безопасный JSON-план.
Отвечай ТОЛЬКО валидным JSON без markdown.
Не выдумывай реквизиты, суммы, даты и юридические условия, которых пользователь не дал. Если точных данных нет, используй понятные заполнители в квадратных скобках: [сумма], [процент], [дата], [этап].

Доступные действия:
1) replace_text: заменить точный фрагмент внутри адресного абзаца или ячейки, сохраняя оформление остальных фрагментов.
   {"type":"replace_text","target_id":"P12","old_text":"старый срок","new_text":"новый срок"}
   Для таблицы target_id имеет вид T2.R4.C3.P1.
2) replace_block: заменить весь адресный абзац/абзац ячейки и при необходимости вставить после него дополнительные абзацы.
   {"type":"replace_block","target_id":"T1.R5.C2.P1","paragraphs":["Новая редакция текста"]}
3) insert_after: вставить абзацы после точного адресного абзаца.
   {"type":"insert_after","target_id":"P18","paragraphs":["Новый пункт 1","Новый пункт 2"]}
4) insert_table_rows: добавить строки в существующую таблицу, копируя оформление соседней строки.
   {"type":"insert_table_rows","table_id":"T2","after_row":4,"rows":[["5","Новая работа","10 дней"]]}
5) append_section: добавить раздел в конец документа.
   {"type":"append_section","heading":"Этапность оплат","paragraphs":["1. ...","2. ..."]}
6) insert_after_heading: резервный вариант для вставки после заголовка/абзаца, который содержит фразу.
   {"type":"insert_after_heading","heading_contains":"оплат","new_heading":"Этапность оплат","paragraphs":["..."]}
7) replace_paragraph_contains: резервный вариант, если точного target_id нет.
   {"type":"replace_paragraph_contains","contains":"старый текст или ключевая фраза","paragraphs":["новая редакция..."]}
8) append_paragraphs: добавить абзацы в конец документа без заголовка.
   {"type":"append_paragraphs","paragraphs":["..."]}
9) add_table: добавить таблицу после адресного абзаца или в конец документа.
   {"type":"add_table","after_target_id":"P25","heading":"График оплат","headers":["Этап","Срок","Размер оплаты"],"rows":[["1","[дата]","[процент]"]]}
10) format_like: привести оформление адресного абзаца к оформлению другого существующего абзаца.
   {"type":"format_like","target_id":"P26","reference_id":"P25"}

Верни объект:
{
  "need_clarification": false,
  "message": "короткое описание, что будет изменено",
  "output_filename": "имя_исправленного_файла.docx",
  "actions": [ ... ]
}

Правила:
- Строки вида [P12], [T2.R4.C3.P1] — точные адреса элементов. Для изменения существующего текста всегда предпочитай target_id.
- Служебные пометки о сокращении текста означают лимит контекста, а не повреждение или обрыв исходного DOCX.
- Документ уже содержит корпоративные стили. Не задавай шрифты, размеры, поля и отступы в JSON: редактор скопирует их из соседнего исходного элемента.
- Для таблиц меняй текст конкретной ячейки через replace_text/replace_block. Не пересоздавай всю таблицу ради нескольких изменений.
- Не заменяй целый абзац, если достаточно заменить короткий фрагмент через replace_text.
- format_like используй только если пользователь явно просит исправить оформление; reference_id должен указывать на правильно оформленный соседний образец.
- Если пользователь просит просто проанализировать документ, верни actions: [] и message с кратким пояснением.
- Если пользователь просит добавить этапность оплат, найди раздел про оплату/расчеты. Если такого раздела не видно — добавь новый раздел в конец: «Этапность оплат».
- Если пользователь просит изменить договор по ТЗ/КП/PDF и такие дополнительные документы переданы ниже, используй их как основание для правок. Не проси прислать ТЗ/КП повторно, если в блоке дополнительных документов уже есть релевантный текст.
- Для договора сохраняй деловой юридический стиль, но не утверждай, что это финальная юридическая редакция.
- Не удаляй большие фрагменты, если пользователь явно не просит.
- Не добавляй предложенные правки отдельным разделом в конец, если в документе видны точные места для замены.
- Не используй макросы, внешние ссылки и произвольный код.
""".strip()
    def make_user_prompt(ref_context: str) -> str:
        return (
            f"Файл: {file_name}\n\n"
            f"Текст DOCX:\n{document_text}\n\n"
            f"Дополнительные документы/контекст (ТЗ, КП, PDF, TXT), если были переданы:\n{ref_context or '[нет]'}\n\n"
            f"Просьба пользователя:\n{request_text}"
        )

    user_prompt = make_user_prompt(reference_context)
    try:
        raw = _call_ai_for_word(system_prompt, user_prompt, max_tokens=7000, temperature=0.0)
    except WordProcessingError as first_error:
        reduced_context = _clip_text_at_boundary(
            reference_context,
            9000,
            "\n... [контекст ТЗ/КП сокращен для повторной попытки]",
        ) if reference_context else ""
        try:
            raw = _call_ai_for_word(system_prompt, make_user_prompt(reduced_context), max_tokens=7000, temperature=0.0)
        except WordProcessingError:
            simple_system_prompt = """
Ты редактируешь DOCX по просьбе пользователя. Ответь только валидным JSON без markdown.
Верни объект:
{
  "need_clarification": false,
  "message": "что будет изменено",
  "output_filename": "edited.docx",
  "actions": [
    {"type":"replace_text","target_id":"P12","old_text":"точный старый текст","new_text":"новый текст"},
    {"type":"replace_block","target_id":"T1.R3.C2.P1","paragraphs":["новый текст ячейки"]},
    {"type":"append_section","heading":"Название раздела","paragraphs":["текст"]}
  ]
}
Используй ТЗ/КП как основание. target_id бери из структуры DOCX.
Не добавляй общий раздел в конец вместо адресных правок существующего текста.
Не проси ТЗ/КП повторно, если они есть в контексте. Не выдумывай неизвестные суммы и реквизиты.
""".strip()
            simple_request = (
                "Сделай минимальный JSON-план правок договора по ТЗ/КП.\n\n"
                f"DOCX:\n{_clip_text_at_boundary(document_text, 45000)}\n\n"
                f"ТЗ/КП:\n{_clip_text_at_boundary(reduced_context, 9000)}\n\n"
                f"Просьба:\n{request_text}\n\n"
                f"Первичная ошибка модели: {first_error}"
            )
            raw = _call_ai_for_word(simple_system_prompt, simple_request, max_tokens=5000, temperature=0.1)
    try:
        patch = _extract_json(raw)
    except WordProcessingError as json_error:
        repair_raw = _call_ai_for_word(
            """
Исправь план редактирования DOCX и верни ТОЛЬКО завершенный валидный JSON без markdown.
Формат: {"need_clarification":false,"message":"...","output_filename":"edited.docx","actions":[]}.
Для точечных правок используй replace_text с target_id, old_text и new_text либо replace_block с target_id и paragraphs.
target_id бери только из структуры документа: P12 или T2.R4.C3.P1.
Не добавляй общий раздел с предложениями в конец вместо реальных правок.
Если надежный адрес правки определить нельзя, верни need_clarification=true, actions=[] и объясни, что уточнить.
""".strip(),
            (
                f"Структура DOCX:\n{_clip_text_at_boundary(document_text, 50000)}\n\n"
                f"ТЗ/КП:\n{_clip_text_at_boundary(reference_context, 10000)}\n\n"
                f"Просьба:\n{request_text}\n\n"
                f"Незавершенный/ошибочный ответ:\n{_clip_text_at_boundary(raw, 9000)}\n\n"
                f"Ошибка JSON: {json_error}"
            ),
            max_tokens=7000,
            temperature=0.0,
        )
        try:
            patch = _extract_json(repair_raw)
        except WordProcessingError as repair_error:
            raise WordProcessingError(
                "ИИ не смог сформировать надежный адресный план правок. "
                "Попробуйте указать конкретный раздел, пункт или фразу для изменения."
            ) from repair_error
    if not isinstance(patch, dict):
        raise WordProcessingError("ИИ вернул не объект JSON")
    patch.setdefault("need_clarification", False)
    patch.setdefault("message", "Готовлю изменения в Word-файле.")
    patch.setdefault("actions", [])
    if not isinstance(patch.get("actions"), list):
        raise WordProcessingError("Поле actions должно быть списком")
    return patch


def _copy_paragraph_format(source: Optional[Paragraph], target: Paragraph) -> None:
    if source is None:
        return
    try:
        existing = target._p.pPr
        if existing is not None:
            target._p.remove(existing)
        if source._p.pPr is not None:
            target._p.insert(0, deepcopy(source._p.pPr))
    except Exception:
        pass
    try:
        if source.style is not None:
            target.style = source.style
    except Exception:
        pass


def _copy_run_format(source: Optional[Paragraph], target_run) -> None:
    if source is None:
        return
    ref_run = next((run for run in source.runs if run.text), None)
    if not ref_run:
        return
    try:
        existing = target_run._r.rPr
        if existing is not None:
            target_run._r.remove(existing)
        if ref_run._r.rPr is not None:
            target_run._r.insert(0, deepcopy(ref_run._r.rPr))
    except Exception:
        pass


def _paragraph_text_nodes(paragraph: Paragraph):
    return list(paragraph._p.iter(qn("w:t")))


def _set_text_node(node, value: str) -> None:
    node.text = value
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"
    if value[:1].isspace() or value[-1:].isspace():
        node.set(xml_space, "preserve")
    elif node.get(xml_space) is not None:
        del node.attrib[xml_space]


def _replace_text_preserving_xml(paragraph: Paragraph, old_text: str, new_text: str) -> bool:
    nodes = _paragraph_text_nodes(paragraph)
    values = [node.text or "" for node in nodes]
    full_text = "".join(values)
    old_text = str(old_text or "")
    if not nodes or not old_text:
        return False

    start = full_text.find(old_text)
    if start < 0:
        start = full_text.lower().find(old_text.lower())
    if start < 0:
        return False
    end = start + len(old_text)

    offsets = []
    cursor = 0
    for value in values:
        offsets.append((cursor, cursor + len(value)))
        cursor += len(value)
    affected = [
        idx
        for idx, (node_start, node_end) in enumerate(offsets)
        if node_end > start and node_start < end
    ]
    if not affected:
        return False

    first_idx = affected[0]
    last_idx = affected[-1]
    first_start, _ = offsets[first_idx]
    _, last_end = offsets[last_idx]
    prefix = values[first_idx][:max(0, start - first_start)]
    suffix_start = max(0, len(values[last_idx]) - (last_end - end))
    suffix = values[last_idx][suffix_start:]

    if first_idx == last_idx:
        _set_text_node(nodes[first_idx], prefix + str(new_text) + suffix)
        return True

    _set_text_node(nodes[first_idx], prefix + str(new_text))
    for idx in affected[1:-1]:
        _set_text_node(nodes[idx], "")
    _set_text_node(nodes[last_idx], suffix)
    return True


def _replace_whole_paragraph_text(paragraph: Paragraph, text: str) -> None:
    nodes = _paragraph_text_nodes(paragraph)
    current = "".join(node.text or "" for node in nodes)
    if nodes and current:
        if _replace_text_preserving_xml(paragraph, current, text):
            return
    _set_paragraph_text_like_runs(paragraph, text, paragraph)


def _set_paragraph_text_like_runs(paragraph: Paragraph, text: str, reference: Optional[Paragraph] = None) -> None:
    """Fallback for a new/empty paragraph that has no reusable XML text nodes."""
    saved_rpr = None
    if reference is not None:
        ref_run = next((run for run in reference.runs if run.text), None)
        if ref_run is not None and ref_run._r.rPr is not None:
            saved_rpr = deepcopy(ref_run._r.rPr)
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        if saved_rpr is not None:
            existing = paragraph.runs[0]._r.rPr
            if existing is not None:
                paragraph.runs[0]._r.remove(existing)
            paragraph.runs[0]._r.insert(0, saved_rpr)
        else:
            _copy_run_format(reference, paragraph.runs[0])
    else:
        run = paragraph.add_run(text)
        if saved_rpr is not None:
            run._r.insert(0, saved_rpr)
        else:
            _copy_run_format(reference, run)


def _set_paragraph_text_like(paragraph: Paragraph, text: str, reference: Optional[Paragraph] = None) -> None:
    nodes = _paragraph_text_nodes(paragraph)
    current = "".join(node.text or "" for node in nodes)
    if current and _replace_text_preserving_xml(paragraph, current, text):
        return
    _set_paragraph_text_like_runs(paragraph, text, reference)


def _last_text_paragraph(doc: Document) -> Optional[Paragraph]:
    for paragraph in reversed(doc.paragraphs):
        if (paragraph.text or "").strip():
            return paragraph
    return doc.paragraphs[-1] if doc.paragraphs else None


def _body_reference_after(anchor: Optional[Paragraph], doc: Document) -> Optional[Paragraph]:
    if anchor is not None:
        paragraphs = list(doc.paragraphs)
        idx = next((i for i, candidate in enumerate(paragraphs) if candidate._p is anchor._p), None)
        if idx is not None:
            for candidate in paragraphs[idx + 1:]:
                if (candidate.text or "").strip():
                    return candidate
        return anchor
    return _last_text_paragraph(doc)


def _insert_paragraph_after(paragraph: Paragraph, text: str = "", style: Optional[str] = None, format_source: Optional[Paragraph] = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    reference = format_source if format_source is not None else paragraph
    _copy_paragraph_format(reference, new_para)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        _set_paragraph_text_like(new_para, text, reference)
    return new_para


def _clear_and_set_paragraph(paragraph: Paragraph, text: str) -> None:
    _replace_whole_paragraph_text(paragraph, text)


def _build_word_targets(doc: Document) -> Tuple[Dict[str, Paragraph], Dict[str, Table]]:
    paragraphs: Dict[str, Paragraph] = {}
    tables: Dict[str, Table] = {}
    paragraph_count = 0
    table_count = 0
    for block in _iter_body_blocks(doc):
        if isinstance(block, Paragraph):
            paragraph_count += 1
            paragraphs[f"P{paragraph_count}"] = block
            continue
        table_count += 1
        table_id = f"T{table_count}"
        tables[table_id] = block
        for r_idx, row in enumerate(block.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                for p_idx, paragraph in enumerate(cell.paragraphs, start=1):
                    paragraphs[f"{table_id}.R{r_idx}.C{c_idx}.P{p_idx}"] = paragraph
    return paragraphs, tables


def _find_heading_reference(doc: Document, level: int = 2) -> Optional[Paragraph]:
    fallback = None
    for paragraph in doc.paragraphs:
        style_name = (paragraph.style.name if paragraph.style is not None else "").lower()
        if "heading" not in style_name and "заголов" not in style_name:
            continue
        fallback = fallback or paragraph
        if str(level) in style_name:
            return paragraph
    return fallback


def _find_paragraph_containing(doc: Document, needle: str) -> Optional[Paragraph]:
    needle = (needle or "").strip().lower()
    if not needle:
        return None
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip().lower()
        if needle in text:
            return paragraph
    # Более мягкий поиск по словам, если точной фразы нет.
    words = [w for w in re.split(r"\W+", needle) if len(w) >= 4]
    if not words:
        return None
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip().lower()
        matches = sum(1 for w in words if w in text)
        if matches >= max(1, min(2, len(words))):
            return paragraph
    return None


def _add_heading(doc: Document, heading: str, level: int = 2, after: Optional[Paragraph] = None) -> Paragraph:
    heading_reference = _find_heading_reference(doc, level)
    if after is not None:
        new_p = OxmlElement("w:p")
        after._p.addnext(new_p)
        paragraph = Paragraph(new_p, after._parent)
        _copy_paragraph_format(heading_reference or after, paragraph)
        _set_paragraph_text_like_runs(paragraph, heading, heading_reference or after)
        return paragraph
    paragraph = doc.add_paragraph()
    _copy_paragraph_format(heading_reference, paragraph)
    if heading_reference is None:
        try:
            paragraph.style = f"Heading {level}"
        except Exception:
            pass
    _set_paragraph_text_like_runs(paragraph, heading, heading_reference)
    return paragraph


def _add_paragraphs(doc: Document, paragraphs: List[str], after: Optional[Paragraph] = None, format_source: Optional[Paragraph] = None) -> Optional[Paragraph]:
    last = after
    if format_source is not None:
        reference = format_source
    elif after is not None:
        reference = after
    else:
        reference = _last_text_paragraph(doc)
    for text in paragraphs or []:
        clean = str(text or "").strip()
        if clean:
            if last is not None:
                last = _insert_paragraph_after(last, clean, format_source=reference if reference is not None else last)
            else:
                p = doc.add_paragraph()
                _copy_paragraph_format(reference, p)
                _set_paragraph_text_like(p, clean, reference)
                last = p
    return last


def _copy_cell_format(source_cell, target_cell) -> None:
    try:
        existing = target_cell._tc.tcPr
        if existing is not None:
            target_cell._tc.remove(existing)
        if source_cell._tc.tcPr is not None:
            target_cell._tc.insert(0, deepcopy(source_cell._tc.tcPr))
    except Exception:
        pass


def _set_cell_text_like(target_cell, text: str, source_cell=None) -> None:
    target_paragraph = target_cell.paragraphs[0]
    source_paragraph = source_cell.paragraphs[0] if source_cell is not None and source_cell.paragraphs else None
    _copy_paragraph_format(source_paragraph, target_paragraph)
    _set_paragraph_text_like(target_paragraph, str(text or ""), source_paragraph)


def _copy_table_format(source: Optional[Table], target: Table) -> None:
    if source is None:
        return
    try:
        existing = target._tbl.tblPr
        if existing is not None:
            target._tbl.remove(existing)
        if source._tbl.tblPr is not None:
            target._tbl.insert(0, deepcopy(source._tbl.tblPr))
    except Exception:
        pass
    if len(source.columns) != len(target.columns):
        return
    try:
        existing_grid = target._tbl.tblGrid
        if existing_grid is not None:
            target._tbl.remove(existing_grid)
        if source._tbl.tblGrid is not None:
            insert_at = 1 if target._tbl.tblPr is not None else 0
            target._tbl.insert(insert_at, deepcopy(source._tbl.tblGrid))
    except Exception:
        pass


def _add_table(doc: Document, heading: Optional[str], headers: List[str], rows: List[List[Any]], after: Optional[Paragraph] = None) -> None:
    if heading:
        after = _add_heading(doc, str(heading), level=2, after=after)
    headers = [str(h or "") for h in (headers or [])]
    rows = rows or []
    if not headers and rows:
        headers = [f"Колонка {i+1}" for i in range(max(len(r) for r in rows))]
    if not headers:
        return
    reference_table = doc.tables[-1] if doc.tables else None
    table_style = reference_table.style if reference_table is not None else "Table Grid"
    table = doc.add_table(rows=1, cols=len(headers))
    _copy_table_format(reference_table, table)
    try:
        table.style = table_style
    except Exception:
        table.style = "Table Grid"
    for idx, header in enumerate(headers):
        source_cell = None
        if reference_table is not None and reference_table.rows:
            source_cell = reference_table.rows[0].cells[min(idx, len(reference_table.rows[0].cells) - 1)]
            _copy_cell_format(source_cell, table.rows[0].cells[idx])
        _set_cell_text_like(table.rows[0].cells[idx], header, source_cell)
    for row in rows[:100]:
        cells = table.add_row().cells
        for idx in range(len(headers)):
            source_cell = None
            if reference_table is not None and reference_table.rows:
                source_row = reference_table.rows[-1]
                source_cell = source_row.cells[min(idx, len(source_row.cells) - 1)]
                _copy_cell_format(source_cell, cells[idx])
            value = row[idx] if idx < len(row) and row[idx] is not None else ""
            _set_cell_text_like(cells[idx], str(value), source_cell)
    if after is not None:
        after._p.addnext(table._tbl)


def _insert_table_rows(table: Table, after_row: int, rows: List[List[Any]]) -> int:
    if not rows or not table.rows:
        return 0
    insert_after = max(1, min(int(after_row or len(table.rows)), len(table.rows)))
    anchor_row = table.rows[insert_after - 1]
    reference_row = anchor_row
    inserted = 0
    for values in rows[:100]:
        new_row = table.add_row()
        try:
            if reference_row._tr.trPr is not None:
                existing = new_row._tr.trPr
                if existing is not None:
                    new_row._tr.remove(existing)
                new_row._tr.insert(0, deepcopy(reference_row._tr.trPr))
        except Exception:
            pass
        for idx, cell in enumerate(new_row.cells):
            source_cell = reference_row.cells[min(idx, len(reference_row.cells) - 1)]
            _copy_cell_format(source_cell, cell)
            value = values[idx] if idx < len(values) and values[idx] is not None else ""
            _set_cell_text_like(cell, str(value), source_cell)
        reference_row._tr.addnext(new_row._tr)
        reference_row = new_row
        inserted += 1
    return inserted


def apply_word_patch(input_path: str, output_path: str, patch: Dict[str, Any]) -> List[str]:
    doc = Document(input_path)
    changes: List[str] = []
    paragraph_targets, table_targets = _build_word_targets(doc)
    actions = patch.get("actions", [])
    if len(actions) > 80:
        raise WordProcessingError("План содержит слишком много правок за один проход. Разделите задачу на части.")

    for action in actions:
        if not isinstance(action, dict):
            raise WordProcessingError("ИИ вернул некорректное действие в плане Word-правок.")
        action_type = action.get("type")

        if action_type == "replace_text":
            target_id = str(action.get("target_id") or "").upper()
            target = paragraph_targets.get(target_id)
            old_text = str(action.get("old_text") or "")
            new_text = str(action.get("new_text") or "")
            if target is None:
                raise WordProcessingError(f"Не найден адрес правки {target_id}.")
            if not _replace_text_preserving_xml(target, old_text, new_text):
                raise WordProcessingError(
                    f"В элементе {target_id} не найден точный текст для замены: «{old_text[:120]}»."
                )
            changes.append(f"Точечно изменён текст в {target_id}")

        elif action_type == "replace_block":
            target_id = str(action.get("target_id") or "").upper()
            target = paragraph_targets.get(target_id)
            paragraphs = [str(p) for p in action.get("paragraphs", []) if str(p or "").strip()]
            if target is None:
                raise WordProcessingError(f"Не найден адрес правки {target_id}.")
            if not paragraphs:
                raise WordProcessingError(f"Для замены {target_id} не передан новый текст.")
            _clear_and_set_paragraph(target, paragraphs[0])
            last = target
            for text in paragraphs[1:]:
                last = _insert_paragraph_after(last, text, format_source=target)
            changes.append(f"Заменён блок {target_id} с сохранением его стиля")

        elif action_type == "insert_after":
            target_id = str(action.get("target_id") or "").upper()
            target = paragraph_targets.get(target_id)
            paragraphs = [str(p) for p in action.get("paragraphs", []) if str(p or "").strip()]
            if target is None:
                raise WordProcessingError(f"Не найден адрес вставки {target_id}.")
            if not paragraphs:
                raise WordProcessingError(f"Для вставки после {target_id} не передан текст.")
            body_reference = _body_reference_after(target, doc)
            _add_paragraphs(doc, paragraphs, after=target, format_source=body_reference or target)
            changes.append(f"После {target_id} добавлено абзацев: {len(paragraphs)}")

        elif action_type == "format_like":
            target_id = str(action.get("target_id") or "").upper()
            reference_id = str(action.get("reference_id") or "").upper()
            target = paragraph_targets.get(target_id)
            reference = paragraph_targets.get(reference_id)
            if target is None or reference is None:
                raise WordProcessingError(
                    f"Не найден адрес форматирования: target={target_id}, reference={reference_id}."
                )
            _copy_paragraph_format(reference, target)
            for run in target.runs:
                _copy_run_format(reference, run)
            changes.append(f"Оформление {target_id} приведено по образцу {reference_id}")

        elif action_type == "insert_table_rows":
            table_id = str(action.get("table_id") or "").upper()
            table = table_targets.get(table_id)
            if table is None:
                raise WordProcessingError(f"Не найдена таблица {table_id}.")
            rows = action.get("rows") or []
            inserted = _insert_table_rows(table, int(action.get("after_row") or len(table.rows)), rows)
            if not inserted:
                raise WordProcessingError(f"Для таблицы {table_id} не переданы строки.")
            changes.append(f"В {table_id} добавлено строк: {inserted}")

        elif action_type == "append_section":
            heading = str(action.get("heading") or "Новый раздел")
            body_ref = _last_text_paragraph(doc)
            heading_para = _add_heading(doc, heading, level=2)
            _add_paragraphs(doc, [str(p) for p in action.get("paragraphs", [])], after=heading_para, format_source=body_ref)
            changes.append(f"Добавлен раздел «{heading}»")

        elif action_type == "insert_after_heading":
            needle = str(action.get("heading_contains") or "")
            heading = str(action.get("new_heading") or action.get("heading") or "Новый раздел")
            anchor = _find_paragraph_containing(doc, needle)
            if not anchor:
                raise WordProcessingError(f"Не найден заголовок/фрагмент для вставки: «{needle}».")
            body_ref = _body_reference_after(anchor, doc)
            last = _add_heading(doc, heading, level=2, after=anchor)
            _add_paragraphs(doc, [str(p) for p in action.get("paragraphs", [])], after=last, format_source=body_ref)
            changes.append(f"Раздел «{heading}» вставлен после фрагмента «{needle}»")

        elif action_type == "replace_paragraph_contains":
            needle = str(action.get("contains") or "")
            paragraphs = [str(p) for p in action.get("paragraphs", []) if str(p or "").strip()]
            if not paragraphs:
                continue
            target = _find_paragraph_containing(doc, needle)
            if target is not None:
                _clear_and_set_paragraph(target, paragraphs[0])
                last = target
                for text in paragraphs[1:]:
                    last = _insert_paragraph_after(last, text, format_source=target)
                changes.append(f"Заменён абзац, содержащий «{needle}»")
            else:
                raise WordProcessingError(f"Не найден фрагмент для замены: «{needle}».")

        elif action_type == "append_paragraphs":
            paragraphs = [str(p) for p in action.get("paragraphs", [])]
            _add_paragraphs(doc, paragraphs, format_source=_last_text_paragraph(doc))
            changes.append(f"Добавлены абзацы: {len([p for p in paragraphs if p.strip()])}")

        elif action_type == "add_table":
            after_target_id = str(action.get("after_target_id") or "").upper()
            after = paragraph_targets.get(after_target_id) if after_target_id else None
            if after_target_id and after is None:
                raise WordProcessingError(f"Не найден адрес для вставки таблицы {after_target_id}.")
            _add_table(
                doc,
                action.get("heading"),
                action.get("headers") or [],
                action.get("rows") or [],
                after=after,
            )
            changes.append(f"Добавлена таблица «{action.get('heading') or 'без заголовка'}»")

        else:
            raise WordProcessingError(f"ИИ предложил неподдерживаемое действие Word: {action_type}.")

    if not changes:
        raise WordProcessingError("План Word-правок не содержит применимых изменений.")
    doc.save(output_path)
    return changes


def _section_signature(doc: Document) -> List[Tuple[Any, ...]]:
    fields = (
        "page_width",
        "page_height",
        "top_margin",
        "right_margin",
        "bottom_margin",
        "left_margin",
        "gutter",
        "header_distance",
        "footer_distance",
        "orientation",
        "start_type",
    )
    return [
        tuple(str(getattr(section, field, None)) for field in fields)
        for section in doc.sections
    ]


def _zip_entries(path: str) -> Dict[str, bytes]:
    with zipfile.ZipFile(path, "r") as archive:
        bad_file = archive.testzip()
        if bad_file:
            raise WordProcessingError(f"Повреждена внутренняя часть DOCX: {bad_file}")
        return {name: archive.read(name) for name in archive.namelist()}


def _asset_hashes(entries: Dict[str, bytes]) -> List[Tuple[str, str]]:
    prefixes = ("word/media/", "word/embeddings/", "word/charts/", "word/diagrams/")
    return sorted(
        (name, hashlib.sha256(data).hexdigest())
        for name, data in entries.items()
        if name.startswith(prefixes)
    )


def _xml_text_for_prefixes(entries: Dict[str, bytes], prefixes: Tuple[str, ...]) -> Dict[str, str]:
    from xml.etree import ElementTree

    result = {}
    for name, data in entries.items():
        if not any(name.startswith(prefix) for prefix in prefixes) or not name.endswith(".xml"):
            continue
        try:
            root = ElementTree.fromstring(data)
            text_parts = [
                node.text or ""
                for node in root.iter()
                if node.tag.endswith("}t") or node.tag.endswith("}instrText")
            ]
            result[name] = "".join(text_parts)
        except Exception:
            result[name] = hashlib.sha256(data).hexdigest()
    return result


def _xml_feature_counts(entries: Dict[str, bytes]) -> Dict[str, int]:
    document_xml = entries.get("word/document.xml", b"")
    return {
        "fields": document_xml.count(b"<w:fldChar") + document_xml.count(b"<w:instrText"),
        "bookmarks": document_xml.count(b"<w:bookmarkStart"),
        "drawings": document_xml.count(b"<w:drawing") + document_xml.count(b"<w:pict"),
        "content_controls": document_xml.count(b"<w:sdt"),
        "tables": document_xml.count(b"<w:tbl"),
    }


def _external_relationships(entries: Dict[str, bytes]) -> List[Tuple[str, str]]:
    from xml.etree import ElementTree

    relationships = []
    for name, data in entries.items():
        if not name.endswith(".rels"):
            continue
        try:
            root = ElementTree.fromstring(data)
        except Exception:
            continue
        for node in root:
            if node.attrib.get("TargetMode") == "External":
                relationships.append((node.attrib.get("Type", ""), node.attrib.get("Target", "")))
    return sorted(relationships)


def _word_definition_ids(entries: Dict[str, bytes], part_name: str, element_suffix: str, attribute_suffix: str) -> List[str]:
    from xml.etree import ElementTree

    data = entries.get(part_name)
    if not data:
        return []
    try:
        root = ElementTree.fromstring(data)
    except Exception:
        return []
    values = []
    for node in root.iter():
        if not node.tag.endswith("}" + element_suffix):
            continue
        for name, value in node.attrib.items():
            if name.endswith("}" + attribute_suffix) or name == attribute_suffix:
                values.append(value)
                break
    return sorted(values)


def validate_word_result(input_path: str, output_path: str, patch: Dict[str, Any]) -> List[str]:
    if not os.path.exists(output_path) or os.path.getsize(output_path) < 500:
        raise WordProcessingError("Получившийся DOCX пуст или не был сохранен.")

    original_entries = _zip_entries(input_path)
    result_entries = _zip_entries(output_path)
    original_doc = Document(input_path)
    result_doc = Document(output_path)

    if _section_signature(original_doc) != _section_signature(result_doc):
        raise WordProcessingError("После правки изменились секции, формат страницы или поля документа.")
    if _asset_hashes(original_entries) != _asset_hashes(result_entries):
        raise WordProcessingError("После правки изменились встроенные изображения, диаграммы или вложения.")

    protected_prefixes = (
        "word/header",
        "word/footer",
        "word/footnotes",
        "word/endnotes",
        "word/comments",
    )
    if _xml_text_for_prefixes(original_entries, protected_prefixes) != _xml_text_for_prefixes(
        result_entries, protected_prefixes
    ):
        raise WordProcessingError("После правки неожиданно изменились колонтитулы, сноски или комментарии.")
    if _external_relationships(original_entries) != _external_relationships(result_entries):
        raise WordProcessingError("После правки изменились внешние ссылки документа.")
    if _word_definition_ids(original_entries, "word/styles.xml", "style", "styleId") != _word_definition_ids(
        result_entries, "word/styles.xml", "style", "styleId"
    ):
        raise WordProcessingError("После правки изменился набор стилей Word.")
    if _word_definition_ids(
        original_entries, "word/numbering.xml", "abstractNum", "abstractNumId"
    ) != _word_definition_ids(result_entries, "word/numbering.xml", "abstractNum", "abstractNumId"):
        raise WordProcessingError("После правки изменилась схема нумерации Word.")

    before_features = _xml_feature_counts(original_entries)
    after_features = _xml_feature_counts(result_entries)
    for feature in ("fields", "bookmarks", "drawings", "content_controls"):
        if after_features[feature] < before_features[feature]:
            raise WordProcessingError(f"После правки потеряны элементы Word: {feature}.")

    add_table_count = sum(1 for action in patch.get("actions", []) if action.get("type") == "add_table")
    if after_features["tables"] < before_features["tables"] + add_table_count:
        raise WordProcessingError("После правки потерялась одна или несколько таблиц.")
    if len(result_doc.paragraphs) < len(original_doc.paragraphs):
        raise WordProcessingError("После правки неожиданно уменьшилось число абзацев документа.")

    return [
        "Проверена целостность DOCX",
        "Сохранены исходные секции, формат страниц, колонтитулы, поля и изображения",
    ]


def edit_word_with_ai(input_path: str, file_name: str, request_text: str, user_id: str, reference_context: str = "") -> Tuple[Optional[str], str, List[str]]:
    patch = build_word_patch_with_ai(input_path, file_name, request_text, reference_context)
    if patch.get("need_clarification"):
        return None, patch.get("message") or "Нужно уточнить, что именно изменить в Word-файле.", []

    actions = patch.get("actions", [])
    if not actions:
        answer = analyze_word_with_ai(input_path, file_name, request_text, user_id, reference_context)
        return None, answer, []

    base = Path(file_name).stem or "document"
    out_name = _safe_filename(patch.get("output_filename") or f"{base}_edited_{uuid.uuid4().hex[:6]}.docx")
    # Чтобы не перезаписать исходный файл даже при совпадении имени.
    if not Path(out_name).stem.endswith("_edited") and "edited" not in out_name.lower() and "исправ" not in out_name.lower():
        out_name = _safe_filename(f"{Path(out_name).stem}_edited.docx")
    output_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex[:6]}_{out_name}")
    try:
        changes = apply_word_patch(input_path, output_path, patch)
        changes.extend(validate_word_result(input_path, output_path, patch))
    except Exception:
        if os.path.exists(output_path):
            try:
                os.unlink(output_path)
            except OSError:
                pass
        raise
    message = patch.get("message") or "Готово, внесла изменения в Word-файл."

    add_to_conversation(user_id, "user", f"[Правка Word {file_name}] {request_text}")
    add_to_conversation(user_id, "assistant", message + "\n" + "\n".join(changes))
    return output_path, message, changes


def _build_word_from_spec(spec: Dict[str, Any], output_path: str) -> None:
    doc = Document()
    title = str(spec.get("title") or "Документ").strip()
    if title:
        p = doc.add_heading(title, level=1)
        try:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
    for section in spec.get("sections") or []:
        heading = str(section.get("heading") or "").strip()
        if heading:
            doc.add_heading(heading, level=int(section.get("level") or 2))
        _add_paragraphs(doc, [str(p) for p in section.get("paragraphs", [])])
        table = section.get("table")
        if isinstance(table, dict):
            _add_table(doc, None, table.get("headers") or [], table.get("rows") or [])
    doc.save(output_path)


def _build_template_context_with_ai(template: Dict[str, str], variables: List[str], request_text: str) -> Dict[str, Any]:
    system_prompt = """
Ты заполняешь Word/DOCX-шаблон компании. Ответь ТОЛЬКО валидным JSON без markdown.
Верни объект:
{
  "filename": "имя_файла.docx",
  "message": "короткое описание результата",
  "context": {"variable_name": "значение"}
}
Правила:
- Заполни ВСЕ переменные шаблона из списка.
- Если данных не хватает, ставь аккуратный заполнитель в квадратных скобках, например [указать сумму].
- Для многострочных блоков используй переносы строк внутри строки.
- Пиши на русском деловом языке.
- Не меняй имена переменных.
""".strip()
    user_prompt = (
        f"Шаблон: {template.get('file_name')}\n"
        f"Переменные шаблона:\n{json.dumps(variables, ensure_ascii=False)}\n\n"
        f"Запрос пользователя:\n{request_text}"
    )
    raw = _call_ai_for_word(system_prompt, user_prompt, max_tokens=4500, temperature=0.1)
    spec = _extract_json(raw)
    context = spec.get("context")
    if not isinstance(context, dict):
        raise WordProcessingError("ИИ не вернул context для Word-шаблона.")
    for variable in variables:
        context.setdefault(variable, f"[{variable}]")
    return {
        "filename": _safe_filename(spec.get("filename") or f"{Path(template.get('name') or 'document').stem}_filled.docx"),
        "message": spec.get("message") or f"Готово, заполнила шаблон {template.get('file_name')}.",
        "context": context,
    }


def create_word_from_template_request(request_text: str) -> Optional[Tuple[str, str]]:
    template = _choose_word_template(request_text)
    if not template:
        return None

    variables = _word_template_variables(template["path"])
    if not variables:
        return None

    spec = _build_template_context_with_ai(template, variables, request_text)
    output_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex[:6]}_{spec['filename']}")
    _render_word_template(template["path"], spec["context"], output_path)
    message = (
        f"{spec['message']}\n"
        f"Использован шаблон: {template.get('file_name')}\n"
        f"Заполнено полей: {len(variables)}"
    )
    return output_path, message


def format_word_templates_list() -> str:
    templates = _list_word_templates()
    if not templates:
        return (
            "Пока нет сохраненных Word-шаблонов.\n"
            "Отправьте .docx с подписью: «это шаблон договора» или «это шаблон КП».\n"
            "Внутри шаблона используйте переменные вида {{ client_name }}, {{ contract_subject }}, {{ price }}."
        )
    lines = ["Сохраненные Word-шаблоны:"]
    for idx, template in enumerate(templates, start=1):
        try:
            variables = _word_template_variables(template["path"])
            suffix = f" · полей: {len(variables)}" if variables else " · без переменных"
        except Exception:
            suffix = ""
        lines.append(f"{idx}. {template['file_name']}{suffix}")
    return "\n".join(lines)


def create_word_from_request(request_text: str) -> Tuple[str, str]:
    templated = create_word_from_template_request(request_text)
    if templated:
        return templated

    system_prompt = """
Ты создаешь структуру Word/DOCX-документа по просьбе пользователя. Отвечай ТОЛЬКО валидным JSON без markdown.
Верни объект:
{
  "filename": "имя_файла.docx",
  "message": "короткое описание файла",
  "title": "Заголовок документа",
  "sections": [
    {"heading":"Раздел", "level":2, "paragraphs":["..."], "table":{"headers":["..."],"rows":[["..."]]}}
  ]
}
Правила:
- Русский деловой стиль.
- Если это договор/акт/письмо, используй аккуратные формулировки и заполнители [ ... ] для неизвестных реквизитов, сумм и дат.
- Не больше 12 разделов.
""".strip()
    user_prompt = f"Создай Word/DOCX-документ по просьбе пользователя:\n{request_text}"
    raw = _call_ai_for_word(system_prompt, user_prompt, max_tokens=3500, temperature=0.2)
    spec = _extract_json(raw)
    file_name = _safe_filename(spec.get("filename") or "created_document.docx")
    output_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex[:6]}_{file_name}")
    _build_word_from_spec(spec, output_path)
    return output_path, spec.get("message") or "Готово, создала Word-файл."


async def handle_word_document(update, context) -> None:
    doc_msg = update.message.document
    file_name = doc_msg.file_name or "document.docx"

    if not is_supported_word_file(file_name):
        await update.message.reply_text(
            "Пока поддерживаю редактирование только .docx. Если у вас .doc/.rtf/.odt — откройте файл в Word и сохраните как .docx, затем отправьте снова."
        )
        return

    user_request = (update.message.caption or "").strip()
    file = await context.bot.get_file(doc_msg.file_id)
    tmp_path = None
    out_path = None

    try:
        await update.message.reply_chat_action(action="typing")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        await file.download_to_drive(tmp_path)

        if _looks_like_template_upload(user_request):
            template = remember_word_template(tmp_path, file_name, user_request)
            try:
                variables = _word_template_variables(template["path"])
            except Exception as exc:
                variables = []
                logger.warning("Saved Word template but could not inspect variables: %s", exc)
            await update.message.reply_text(
                f"Сохранила Word-шаблон: {template['file_name']}\n"
                f"Найдено полей для заполнения: {len(variables)}\n\n"
                "Теперь можно просить: «Создай договор в Word по шаблону...» или «Сформируй КП в Word...»."
            )
            return

        user_id = get_dialog_key(update)
        remember_word_file(context, update.effective_chat.id, tmp_path, file_name, user_id)
        remembered_text = document_to_text(
            tmp_path,
            max_paragraphs=500,
            max_tables=30,
            max_table_rows=150,
            max_cell_chars=1400,
        )
        is_edit_request = looks_like_word_edit_request(user_request)
        if not is_edit_request:
            _remember_reference_document(context, user_id, file_name, remembered_text)
        reference_context = _format_reference_documents(_get_reference_documents(context, user_id))

        if not user_request:
            text = remembered_text
            add_to_conversation(user_id, "user", f"[Word {file_name}]\n{text[:5000]}")
            await update.message.reply_text(
                f"Word-файл {file_name} получен и прочитан.\n"
                "Я запомнила этот файл для текущего чата. Следующим сообщением напишите, что сделать, например:\n"
                "• Проанализируй договор и найди спорные места\n"
                "• Добавь раздел «Этапность оплат»\n"
                "• Замени пункт про сроки выполнения работ\n\n"
                "Я пришлю новую .docx-копию, исходный файл не перезаписывается."
            )
            return

        if is_edit_request:
            out_path, message, changes = await asyncio.to_thread(edit_word_with_ai, tmp_path, file_name, user_request, user_id, reference_context)
            if out_path:
                delivery_name = _delivery_filename(out_path)
                remember_word_file(context, update.effective_chat.id, out_path, delivery_name, user_id)
                text_msg = message
                if changes:
                    text_msg += "\n\nЧто изменено:\n" + "\n".join(f"• {change}" for change in changes[:12])
                await update.message.reply_text(text_msg)
                with open(out_path, "rb") as f:
                    await update.message.reply_document(document=f, filename=delivery_name)
            else:
                await update.message.reply_text(message)
        else:
            answer = await asyncio.to_thread(analyze_word_with_ai, tmp_path, file_name, user_request, user_id, reference_context)
            await update.message.reply_text(answer)

    except WordProcessingError as e:
        logger.exception(f"Word AI processing error for {file_name}: {e}")
        await update.message.reply_text(
            "Word-файл прочитан, но не удалось выполнить обработку через ИИ.\n"
            f"Причина: {e}\n\n"
            "Файл не обязательно поврежден. Часто это бывает из-за лимита контекста, сбоя API или некорректного JSON-ответа модели."
        )
    except Exception as e:
        logger.exception(f"Word processing error for {file_name}: {e}")
        await update.message.reply_text(
            "Не удалось обработать Word-файл. Проверьте, что это обычный .docx без пароля и повреждений."
        )
    finally:
        for path in (tmp_path, out_path):
            if path and os.path.exists(path):
                try:
                    os.unlink(path)
                except OSError:
                    pass


async def handle_word_followup_text(update, context, text: str) -> bool:
    dialog_key = get_dialog_key(update)
    info = _get_recent_word_context(context, dialog_key)
    if not info or not _is_word_followup_request(context, text, dialog_key):
        return False

    request_text = (text or "").strip()
    if not request_text:
        return False

    user_id = dialog_key
    input_path = info["path"]
    file_name = info.get("file_name") or "document.docx"
    reference_context = _format_reference_documents(_get_reference_documents(context, dialog_key))
    out_path = None

    try:
        await update.message.reply_chat_action(action="typing")
        if looks_like_word_edit_request(request_text):
            await update.message.reply_text("Поняла. Вношу изменения в последний Word-файл и пришлю новую копию.")
            out_path, message, changes = await asyncio.to_thread(edit_word_with_ai, input_path, file_name, request_text, user_id, reference_context)
            if out_path:
                delivery_name = _delivery_filename(out_path)
                remember_word_file(context, update.effective_chat.id, out_path, delivery_name, dialog_key)
                text_msg = message
                if changes:
                    text_msg += "\n\nЧто изменено:\n" + "\n".join(f"• {change}" for change in changes[:12])
                await update.message.reply_text(text_msg)
                with open(out_path, "rb") as f:
                    await update.message.reply_document(document=f, filename=delivery_name)
            else:
                await update.message.reply_text(message)
        else:
            answer = await asyncio.to_thread(analyze_word_with_ai, input_path, file_name, request_text, user_id, reference_context)
            await update.message.reply_text(answer)
        context.user_data.get("awaiting_word_request_by_dialog", {}).pop(dialog_key, None)
        return True
    except WordProcessingError as e:
        logger.exception(f"Word followup AI processing error for {file_name}: {e}")
        await update.message.reply_text(
            "Последний Word-файл прочитан, но не удалось выполнить обработку через ИИ.\n"
            f"Причина: {e}"
        )
        return True
    except Exception as e:
        logger.exception(f"Word followup processing error for {file_name}: {e}")
        await update.message.reply_text(
            "Не удалось выполнить действие с последним Word-файлом. Попробуйте отправить файл ещё раз с подписью-командой."
        )
        return True
    finally:
        if out_path and os.path.exists(out_path):
            try:
                os.unlink(out_path)
            except OSError:
                pass


async def handle_create_word_text(update, context, text: str) -> bool:
    if _looks_like_template_list_request(text):
        await update.message.reply_text(format_word_templates_list())
        return True

    if not is_create_word_request(text):
        return False

    out_path = None
    try:
        await update.message.reply_chat_action(action="upload_document")
        out_path, message = await asyncio.to_thread(create_word_from_request, text)
        delivery_name = _delivery_filename(out_path)
        remember_word_file(context, update.effective_chat.id, out_path, delivery_name, get_dialog_key(update))
        await update.message.reply_text(message)
        with open(out_path, "rb") as f:
            await update.message.reply_document(document=f, filename=delivery_name)
        return True
    except Exception as e:
        logger.exception(f"Create Word error: {e}")
        await update.message.reply_text("Не удалось создать Word-файл. Попробуйте описать документ проще: тип документа, разделы и ключевые условия.")
        return True
    finally:
        if out_path and os.path.exists(out_path):
            try:
                os.unlink(out_path)
            except OSError:
                pass
